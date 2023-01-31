from django.shortcuts import render,redirect
from Atendimento.forms import LoginForm
from django.views import generic
from .models import plano
from django.db import IntegrityError
from django.template import loader
from django.contrib import messages
from docx import Document
from django.http.response import HttpResponse
from django.views.generic import ListView
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.contrib.auth import authenticate,logout,login
from django.contrib.auth.decorators import login_required
from docxtpl import DocxTemplate
from django.shortcuts import render
from .models import Acesso
from datetime import date
import datetime

# Create your views here.


def view_login(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            user = authenticate(
                request,
                username = form.cleaned_data['username'],
                password = form.cleaned_data['password']

                )
            
            if user is not None:
                login(request,user)
                return redirect('index')
            else:
                messages.error(request,'Credenciais Inválidas')
                return redirect('login')
    else:
        form = LoginForm()
    return render(request, 'Atendimento/templates/login2.html', {'form':form})




class IndexView(generic.ListView):
    model = plano
    template_name = 'Atendimento/templates/index.html'
    def get_queryset(self):
        
        return plano.objects.order_by('-created') 
    def increment_access_count(self):
        today = date.today()
    
        access_count, created = Acesso.objects.get_or_create(date=today)
        access_count.count += 1
        access_count.save()
    def get(self, request, *args, **kwargs):
        self.increment_access_count()
        return super().get(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['today'] = datetime.datetime.now()
        return context



'''class CreateView(generic.ListView):
    model = plano
    template_name = 'Atendimento/templates/criarPlano.html'
    def get_queryset(self):
        return plano.objects.order_by('-created')'''

# classe para criar o pda


class CreatePdaView(CreateView):
    model = plano
    CHOICES = plano.CHOICES
    template_name = 'Atendimento/templates/criarPlano.html'
    fields = ['versao', 'sistema', 'inquilino',  'sub_demanda','demanda',
              'desc_produto', 'n_contrato', 'aditivo', 'ans',
              'necessidades', 'requisitos_LDPA', 'solucoes_impactadas',
              'anexo_contexto', 'telas_sistema', 'cronograma_implantacao','canais','n1','n2','area_solicitante_cliente',
              'obs_solicitante_cliente','area_executora_cliente','obs_executora_cliente',
              'area_solicitante_usuario','obs_solicitante_usuario','area_executora_usuario',
              'obs_executora_usuario','area_solicitante_base','publicacao_solicitante_base','obs_solicitante_base',
              'area_executora_base','publicacao_executora_base','obs_executora_base','area_solicitante_itens',
              'obs_solicitante_itens','area_executora_itens','obs_executora_itens','area_solicitante_cadastro_clientes',
              'area_executora_cadastro_clientes','obs_executora_cliente','obs_executora_cadastro_clientes','obs_solicitante_cadastro_clientes','area_solicitante_cadastro_base',
              'publicacao_executora_cadastro_base','publicacao_solicitante_cadastro_base','obs_solicitante_cadastro_base','area_executora_cadastro_base',
              'obs_executora_cadastro_base','area_responsavel_capacitacao','atendimento','carga_horaria','cronograma','obs_risco',
                'adequacao','capacitacao','responsavel','status','identificacao_area_solicitante','objetivo_plano_atendimento','requisitos_cadastramento_clienteEusuario',
                'abrangencia','Estrategia_de_implan_solucao','Volumetrias','fluxo_de_atendimento','n1Area','n2Area','ano_contrato','riscos_identificados',
                'mitigacoes','impactos','atividadeCronograma', 'AreaCronograma','ResponsavelCronograma', 'DatainiCronograma', 'DatafimCronograma', 'statusCronograma',
                'informacoes_complementares','cadastramento_itens_catalogo']

    success_url = reverse_lazy('index')

    def form_valid(self, form):
        try:
            return super().form_valid(form)
        except IntegrityError as e:
            # Adicione a mensagem de erro ao formulário
            form.add_error(None, str(e))
            return self.form_invalid(form)

    def get_queryset(self):
        return plano.objects.order_by('-created')
    

# função para visualização
@login_required
def ver(request, pda_id):
    pda = plano.objects.filter(id=pda_id)
    pda = pda[0]
    template = loader.get_template('Atendimento/templates/pda.html')
    context = {
        'id': pda.id,
        'sistema': pda.sistema,
        'inquilino': pda.inquilino,
        'desc_produto': pda.desc_produto,
        'CHOICES': pda.CHOICES,
        'CHOICES2': pda.CHOICES2,
        'versao': pda.versao,
        'status': pda.status,
        'created': pda.created,
        'modified': pda.modified,
        'demanda': pda.demanda,
        'sub_demanda': pda.sub_demanda,
        'n_contrato' : pda.n_contrato,
        'aditivo': pda.aditivo,
        'ans': pda.ans,
        'necessidades': pda.necessidades,
        'requisitos_LDPA': pda.requisitos_LDPA,
        'solucoes_impactadas': pda.solucoes_impactadas,
        'anexo_contexto': pda.anexo_contexto,
        'telas_sistema' : pda.telas_sistema,
        'cronograma_implantacao' : pda.cronograma_implantacao,
        'canais' : pda.canais,
        'n1': pda.n1,
        'n2' : pda.n2,
        'adequacao': pda.adequacao,
        'area_solicitante_cliente': pda.area_solicitante_cliente,
        'obs_solicitante_cliente': pda.obs_solicitante_cliente,
        'area_executora_cliente': pda.area_executora_cliente,
        'obs_executora_cliente': pda.obs_executora_cliente,
        'area_solicitante_usuario': pda.area_solicitante_usuario,
        'obs_solicitante_usuario': pda.obs_solicitante_usuario,
        'area_executora_usuario': pda.area_executora_usuario,
        'obs_executora_usuario': pda.obs_executora_usuario,
        'area_solicitante_base': pda.area_solicitante_base,
        'publicacao_solicitante_base' : pda.publicacao_solicitante_base,
        'obs_solicitante_base': pda.obs_solicitante_base,
        'area_executora_base': pda.area_executora_base,
        'publicacao_executora_base': pda.publicacao_executora_base,
        'obs_executora_base': pda.obs_executora_base,
        'area_solicitante_itens': pda.area_solicitante_itens,
        'obs_solicitante_itens': pda.obs_solicitante_itens,
        'area_executora_itens': pda.area_executora_itens,
        'area_solicitante_cadastro_clientes': pda.area_solicitante_cadastro_clientes,
        'area_executora_cadastro_clientes': pda.area_executora_cadastro_clientes,
        'obs_executora_cadastro_clientes': pda.obs_executora_cadastro_clientes,
        'area_solicitante_cadastro_base': pda.area_solicitante_cadastro_base,
        'publicacao_solicitante_cadastro_base': pda.publicacao_solicitante_cadastro_base,
        'obs_solicitante_cadastro_base':pda.obs_solicitante_cadastro_base,
        'area_executora_cadastro_base':pda.area_executora_cadastro_base,
        'publicacao_executora_cadastro_base': pda.publicacao_executora_cadastro_base,
        'obs_executora_cadastro_base': pda.obs_executora_cadastro_base,
        'capacitacao': pda.capacitacao,
        'area_responsavel_capacitacao': pda.area_responsavel_capacitacao,
        'atendimento': pda.atendimento,
        'carga_horaria': pda.carga_horaria,
        'cronograma': pda.cronograma,
        'obs_risco': pda.obs_risco,
        'responsavel': pda.responsavel,
        'obs_executora_itens':pda.obs_executora_itens,
        'obs_solicitante_cadastro_clientes':pda.obs_solicitante_cadastro_clientes,
        'identificacao_area_solicitante':pda.identificacao_area_solicitante,
        'objetivo_plano_atendimento':pda.objetivo_plano_atendimento,
        'requisitos_cadastramento_clienteEusuario':pda.requisitos_cadastramento_clienteEusuario,
        'abrangencia':pda.abrangencia,
        'Estrategia_de_implan_solucao':pda.Estrategia_de_implan_solucao,
        'Volumetrias':pda.Volumetrias,
        'fluxo_de_atendimento':pda.fluxo_de_atendimento,
        'n1Area':pda.n1Area,
        'n2Area':pda.n2Area,
        'ano_contrato':pda.ano_contrato,
        'riscos_identificados':pda.riscos_identificados,
        'impactos':pda.impactos,
        'mitigacoes':pda.mitigacoes,

        'atividadeCronograma':pda.atividadeCronograma,
        'AreaCronograma':pda.AreaCronograma,
        'ResponsavelCronograma':pda.ResponsavelCronograma,
        'DatainiCronograma':pda.DatainiCronograma,
        'DatafimCronograma':pda.DatafimCronograma,
        'statusCronograma':pda.statusCronograma,
        'informacoes_complementares':pda.informacoes_complementares,
        'cadastramento_itens_catalogo':pda.cadastramento_itens_catalogo

        



    }
    return HttpResponse(template.render(context, request))

# classe para atualizar o pda
class UpdatePda(UpdateView):
    model = plano
    template_name = 'Atendimento/templates/atualizar.html'
    fields = ['versao','status', 'sistema','inquilino','demanda','sub_demanda',
              'desc_produto','n_contrato','aditivo','ans',
              'necessidades','requisitos_LDPA','solucoes_impactadas',
              'anexo_contexto','telas_sistema','cronograma_implantacao',
              'canais','n1','n2',  'area_solicitante_cliente',
              'obs_solicitante_cliente','area_executora_cliente','obs_executora_cliente',
              'area_solicitante_usuario','obs_solicitante_usuario','area_executora_usuario',
              'obs_executora_usuario','area_solicitante_base','obs_solicitante_base',
              'area_executora_base','publicacao_executora_base','obs_executora_base','area_solicitante_itens',
              'obs_solicitante_itens','area_executora_itens','obs_executora_itens','area_solicitante_cadastro_clientes',
              'area_executora_cadastro_clientes','obs_executora_cliente','area_solicitante_cadastro_base',
              'publicacao_executora_cadastro_base','obs_solicitante_cadastro_base','area_executora_cadastro_base',
              'publicacao_executora_cadastro_base','obs_executora_cadastro_base','capacitacao',
              'area_responsavel_capacitacao','atendimento','carga_horaria','cronograma','obs_risco','adequacao'
              ,'responsavel','obs_solicitante_cadastro_clientes','publicacao_solicitante_base','identificacao_area_solicitante','objetivo_plano_atendimento','requisitos_cadastramento_clienteEusuario',
              'abrangencia','Estrategia_de_implan_solucao','Volumetrias','fluxo_de_atendimento','n1Area','n2Area','ano_contrato',
              'riscos_identificados','impactos','mitigacoes','atividadeCronograma', 'AreaCronograma','ResponsavelCronograma', 'DatainiCronograma', 'DatafimCronograma', 'statusCronograma',
              'informacoes_complementares','cadastramento_itens_catalogo']

    success_url = reverse_lazy('index')
    def form_valid(self, form):
        try:
            return super().form_valid(form)
        except IntegrityError as e:
            # Adicione a mensagem de erro ao formulário
            form.add_error(None, str(e))
            return self.form_invalid(form)





class DeletePda(DeleteView):
    model = plano
    abstract = True

    template_name = 'Atendimento/templates/deletado.html'
    success_url = reverse_lazy('index')


    
@login_required
def desativar(request, pk):
    template = loader.get_template('Atendimento/templates/desativado.html')
    pda = plano.objects.filter(id=pk)
    pda = pda[0]
    context = {
        'sistema': pda.sistema,
        'inquilino': pda.inquilino,
        
        }
    if (request.method == 'POST'):
        plano.objects.filter(id=pk).update(status = '4')
        print('Desativado')
        reverse_lazy('index')
        messages.info(request, 'Plano concluido com sucesso!')    
    else:
        print('(Erro!) Não concluído!.')

    

    return HttpResponse(template.render( context,request))



'''def Mudando(request, pk):
    template = loader.get_template('Atendimento/templates/desativado.html')
    context = {
        'sistema': plano.sistema,
        'inquilino': plano.inquilino,
        'desc_produto': plano.desc_produto,
        
        
        }
    
    plano.objects.filter(id=pk).update(status = '4')
    print('Desativado')
   

    return HttpResponse(template.render( context,request))'''



#plano.objects.filter(id=pk).update(status = '4')


    
    
  


'''def criaWord(request,pk):
    pda_ed = plano.objects.filter(id=pk)
    pda_ed = pda_ed[0]

    context = {
        'id': pda_ed.id,
        'sistema': pda_ed.sistema
    }
    document = Document()
    document.add_heading('sistema', 0)
    document.save('demo.docx')

    template = loader.get_template('Atendimento/templates/editar.html')

    return HttpResponse(template.render(context, request))'''

    
'''@login_required
def download_world(request,pk):
    pda_ed = plano.objects.filter(id=pk)
    pda_ed = pda_ed[0]


    document = Document()

    document.add_heading('Proposta de Atendimento', 0)
    document.add_paragraph(pda_ed.sistema)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=proposta.docx'
    document.save(response)
    return response
'''
@login_required
def download_world(request,pk):
    template = DocxTemplate("Template_PDA.docx")
    spa = plano.objects.get(id=pk)
    
    context = {
        'id': spa.id,
        'sistema': spa.sistema,
        'inquilino': spa.inquilino,
        'desc_produto': spa.desc_produto,
        'CHOICES': spa.CHOICES,
        'CHOICES2': spa.CHOICES2,
        'versao': spa.versao,
        'status': spa.status,
        'created': spa.created,
        'modified': spa.modified,
        'demanda': spa.demanda,
        'sub_demanda': spa.sub_demanda,
        'n_contrato' : spa.n_contrato,
        'aditivo': spa.aditivo,
        'ans': spa.ans,
        'necessidades': spa.necessidades,
        'requisitos_LDPA': spa.requisitos_LDPA,
        'solucoes_impactadas': spa.solucoes_impactadas,
        'anexo_contexto': spa.anexo_contexto,
        'telas_sistema' : spa.telas_sistema,
        'cronograma_implantacao' : spa.cronograma_implantacao,
        'canais' : spa.canais,
        'n1': spa.n1,
        'n2' : spa.n2,
        'adequacao': spa.adequacao,
        'area_solicitante_cliente': spa.area_solicitante_cliente,
        'obs_solicitante_cliente': spa.obs_solicitante_cliente,
        'area_executora_cliente': spa.area_executora_cliente,
        'obs_executora_cliente': spa.obs_executora_cliente,
        'area_solicitante_usuario': spa.area_solicitante_usuario,
        'obs_solicitante_usuario': spa.obs_solicitante_usuario,
        'area_executora_usuario': spa.area_executora_usuario,
        'obs_executora_usuario': spa.obs_executora_usuario,
        'area_solicitante_base': spa.area_solicitante_base,
        'publicacao_solicitante_base' : spa.publicacao_solicitante_base,
        'obs_solicitante_base': spa.obs_solicitante_base,
        'area_executora_base': spa.area_executora_base,
        'publicacao_executora_base': spa.publicacao_executora_base,
        'obs_executora_base': spa.obs_executora_base,
        'area_solicitante_itens': spa.area_solicitante_itens,
        'obs_solicitante_itens': spa.obs_solicitante_itens,
        'area_executora_itens': spa.area_executora_itens,
        'area_solicitante_cadastro_clientes': spa.area_solicitante_cadastro_clientes,
        'area_executora_cadastro_clientes': spa.area_executora_cadastro_clientes,
        'obs_executora_cadastro_clientes': spa.obs_executora_cadastro_clientes,
        'area_solicitante_cadastro_base': spa.area_solicitante_cadastro_base,
        'publicacao_solicitante_cadastro_base': spa.publicacao_solicitante_cadastro_base,
        'obs_solicitante_cadastro_base':spa.obs_solicitante_cadastro_base,
        'area_executora_cadastro_base':spa.area_executora_cadastro_base,
        'publicacao_executora_cadastro_base': spa.publicacao_executora_cadastro_base,
        'obs_executora_cadastro_base': spa.obs_executora_cadastro_base,
        'capacitacao': spa.capacitacao,
        'area_responsavel_capacitacao': spa.area_responsavel_capacitacao,
        'atendimento': spa.atendimento,
        'carga_horaria': spa.carga_horaria,
        'cronograma': spa.cronograma,
        'obs_risco': spa.obs_risco,
        'responsavel': spa.responsavel,
        'obs_executora_itens':spa.obs_executora_itens,
        'obs_solicitante_cadastro_clientes':spa.obs_solicitante_cadastro_clientes,
        'identificacao_area_solicitante':spa.identificacao_area_solicitante,
        'objetivo_plano_atendimento':spa.objetivo_plano_atendimento,
        'requisitos_cadastramento_clienteEusuario':spa.requisitos_cadastramento_clienteEusuario,
        'abrangencia':spa.abrangencia,
        'Estrategia_de_implan_solucao':spa.Estrategia_de_implan_solucao,
        'Volumetrias':spa.Volumetrias,
        'fluxo_de_atendimento':spa.fluxo_de_atendimento,
        'n1Area':spa.n1Area,
        'n2Area':spa.n2Area,
        'ano_contrato':spa.ano_contrato,
        'riscos_identificados':spa.riscos_identificados,
        'impactos':spa.impactos,
        'mitigacoes':spa.mitigacoes,
        'atividadeCronograma':spa.atividadeCronograma,
        'AreaCronograma':spa.AreaCronograma,
        'ResponsavelCronograma':spa.ResponsavelCronograma,
        'DatainiCronograma':spa.DatainiCronograma,
        'DatafimCronograma':spa.DatafimCronograma,
        'statusCronograma':spa.statusCronograma,
        'informacoes_complementares':spa.informacoes_complementares,
        'cadastramento_itens_catalogo':spa.cadastramento_itens_catalogo
    }
    template.render(context)
    

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=documento_preenchido.docx'
    
    template.save(response)

    return response



# duplicar plano

def duplicate_model(request, id):
    # Recupera o objeto a ser duplicado
    plano_original = plano.objects.get(id=id)

    # Cria um novo objeto com os mesmos campos
    planoDuplicado = plano_original(
        versao=plano_original.versao,
        
        # ...
    )
    planoDuplicado.save()

    # Renderiza a página atual
    return render(request, 'Atendimento/templates/pda.html', {'planoDuplicado': planoDuplicado})
